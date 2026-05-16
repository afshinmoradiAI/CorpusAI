"""Convert assembled markdown into a .docx byte stream.

The paper / grant / thesis format is well-constrained (we wrote the
assembler ourselves), so a small line-by-line walker handles every shape
we emit:

    # Title              -> Heading 0
    ## Section           -> Heading 1
    ### Subsection       -> Heading 2
    ![alt](url)          -> embedded image (Thesis only — resolved via figure_store)
    [N] citation         -> numbered reference paragraph
    plain line           -> body paragraph
    blank line           -> paragraph break (already implied)

Output uses Times New Roman throughout (body 12pt, headings 14/13/12pt).
"""

from __future__ import annotations

import io
import re

from docx import Document  # type: ignore[import-untyped]
from docx.shared import Inches, Pt

from app.services.figure_store import path_for

_FONT_NAME = "Times New Roman"
_BODY_SIZE = Pt(12)
_HEADING_SIZES = {0: Pt(16), 1: Pt(14), 2: Pt(13)}
_MAX_IMAGE_WIDTH_INCHES = 5.5  # fits A4 with 2.5cm margins

# Markdown image syntax for figures the workflow emits.
# Pattern: ![alt text](/api/thesis/figure/<figure_id>)
_FIGURE_LINE_RE = re.compile(r"^!\[[^\]]*\]\(/api/thesis/figure/([a-zA-Z0-9_\-]+)\)\s*$")


def _apply_font(paragraph, size: Pt) -> None:
    for run in paragraph.runs:
        run.font.name = _FONT_NAME
        run.font.size = size


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = _FONT_NAME
    style.font.size = _BODY_SIZE


def paper_to_docx(markdown: str) -> bytes:
    doc = Document()
    _set_default_style(doc)

    for line in markdown.splitlines():
        stripped = line.rstrip()
        if not stripped:
            continue

        figure_match = _FIGURE_LINE_RE.match(stripped)
        if figure_match:
            figure_id = figure_match.group(1)
            _add_figure(doc, figure_id)
            continue

        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:].strip(), level=0)
            _apply_font(h, _HEADING_SIZES[0])
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:].strip(), level=1)
            _apply_font(h, _HEADING_SIZES[1])
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:].strip(), level=2)
            _apply_font(h, _HEADING_SIZES[2])
        else:
            p = doc.add_paragraph(stripped)
            _apply_font(p, _BODY_SIZE)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_figure(doc: Document, figure_id: str) -> None:
    """Embed a figure by ID. Silently skips missing files (graceful fallback)."""
    path = path_for(figure_id)
    if path is None:
        # Leave a visible marker rather than failing the whole export.
        p = doc.add_paragraph(f"[Figure unavailable: {figure_id}]")
        _apply_font(p, _BODY_SIZE)
        return
    try:
        doc.add_picture(str(path), width=Inches(_MAX_IMAGE_WIDTH_INCHES))
    except Exception:  # noqa: BLE001 — never crash an export on a single image
        p = doc.add_paragraph(f"[Figure could not be embedded: {figure_id}]")
        _apply_font(p, _BODY_SIZE)
