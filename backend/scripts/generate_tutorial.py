"""Generate CorpusAI_Tutorial.docx — a guided tour of the app.

Run from the backend/ directory:

    uv run python scripts/generate_tutorial.py

The output lands at the repo root as CorpusAI_Tutorial.docx.

This script intentionally has no production dependencies — only python-docx,
which is already in the project. Re-run it whenever the app changes and you
want a refreshed tutorial.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[2] / "CorpusAI_Tutorial.docx"

CODE_FONT = "Consolas"
CODE_SIZE = Pt(9.5)
CODE_FILL = "F4F4F4"
INLINE_FILL = "EFEFEF"
MUTED = RGBColor(0x55, 0x55, 0x55)


# ---------- low-level helpers ----------


def _shade(paragraph, fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code_block(doc: Document, code: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(code)
    run.font.name = CODE_FONT
    run.font.size = CODE_SIZE
    _shade(para, CODE_FILL)


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def kvp(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(value)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = MUTED
    run.font.size = Pt(9)


# ---------- content ----------


def build_document() -> Document:
    doc = Document()

    # ---- cover ----
    title = doc.add_heading("CorpusAI", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Build & Use Tutorial — a multi-agent biology research assistant")
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    body(
        doc,
        "This document has two halves. Part A teaches you how to use the app. "
        "Part B teaches you how it was built, file by file, so you can extend it.",
    )
    doc.add_paragraph()

    # ============================================================
    # PART A — USING THE APP
    # ============================================================
    doc.add_heading("Part A — Using CorpusAI", level=1)

    doc.add_heading("What it does", level=2)
    body(
        doc,
        "CorpusAI is a web application that helps biology researchers in two ways:",
    )
    bullet(
        doc,
        "Explore mode — you type a topic; the app searches the literature, "
        "identifies a research gap, and proposes one paragraph each of a "
        "novel idea, a hypothetical method, and the aim & importance.",
    )
    bullet(
        doc,
        "Write mode — you upload reference PDFs and CorpusAI drafts a full "
        "academic paper (Abstract, Introduction, Methods, Results, Discussion) "
        "with three peer reviews (biology, statistics, gap) plus a synthesised "
        "revision list.",
    )
    body(
        doc,
        "Behind the curtain, 18 specialised LLM agents are wired into pipelines "
        "by an orchestrator. The frontend streams progress over Server-Sent "
        "Events so you can watch each step happen live.",
    )

    doc.add_heading("Prerequisites", level=2)
    bullet(doc, "Python 3.11 (uv will download it automatically if missing)")
    bullet(doc, "Node 20+ and npm")
    bullet(doc, "An Anthropic API key — set ANTHROPIC_API_KEY in backend/.env")

    doc.add_heading("Running locally", level=2)
    body(doc, "Open two terminals.")
    body(doc, "Terminal 1 — backend:")
    code_block(
        doc,
        "cd backend\n"
        "cp .env.example .env          # then edit .env to add your Anthropic key\n"
        "uv sync                       # installs dependencies\n"
        "uv run uvicorn app.main:app --reload\n"
        "# server: http://localhost:8000",
    )
    body(doc, "Terminal 2 — frontend:")
    code_block(
        doc,
        "cd frontend\n"
        "npm install\n"
        "npm run dev\n"
        "# UI: http://localhost:3000",
    )
    body(
        doc,
        "Open http://localhost:3000. You will see two tabs: Explore and Write.",
    )

    doc.add_heading("Using Explore mode", level=2)
    numbered(doc, "Click the Explore tab.")
    numbered(
        doc,
        "Enter a biology topic. Be specific — \"Mitochondrial dynamics in T-cell "
        "exhaustion\" produces better results than just \"T cells\".",
    )
    numbered(
        doc,
        "Click Generate idea. The left sidebar shows a step list — each step "
        "switches to amber while running and green when done.",
    )
    numbered(
        doc,
        "When the pipeline finishes, the right panel shows four sections: "
        "Research gap, Idea, Hypothetical method, and Aim & importance, "
        "followed by the references list.",
    )
    body(
        doc,
        "If literature search returns no papers (rare on common topics), "
        "the pipeline emits an error event and stops gracefully.",
    )

    doc.add_heading("Using Write mode", level=2)
    numbered(doc, "Click the Write tab.")
    numbered(
        doc,
        "Upload reference PDFs (up to 30 files, 25 MB each). The backend parses "
        "every page and builds a BM25 retrieval index in memory.",
    )
    numbered(doc, "Enter the paper topic.")
    numbered(
        doc,
        "Choose which sections to generate. Tick all five for a full paper, or "
        "any subset — for example, just Methods if you only need that section.",
    )
    numbered(
        doc,
        "(Optional) expand Optional context to paste raw experimental results "
        "or notes for the writer agents.",
    )
    numbered(
        doc,
        "Click Generate paper. Each section streams in: amber while writing, "
        "green when done. The references are formatted, the markdown is "
        "assembled, and the three reviewers run in parallel.",
    )
    numbered(
        doc,
        "When complete, the assembled paper is displayed at the top, with "
        "the four-tab Peer review panel below (Synthesis, Biology, Statistics, "
        "Gap).",
    )
    numbered(
        doc,
        "Click Download .docx to save a Word file. The .md button gives you "
        "the same content as plain markdown.",
    )

    doc.add_heading("Reading the peer review", level=2)
    body(
        doc,
        "The review panel has four tabs:",
    )
    kvp(
        doc,
        "Synthesis",
        "An executive summary that merges all three reviews plus a "
        "prioritised list of revisions, ordered by impact.",
    )
    kvp(
        doc,
        "Biology",
        "A senior PI critique — accuracy, missing controls, terminology "
        "errors, over-interpretation of correlative data. Includes an "
        "overall score (1–5) and per-issue severity (critical/major/minor).",
    )
    kvp(
        doc,
        "Statistics",
        "A biostatistician critique — sample size, test choice, multiple "
        "comparisons, effect sizes, reproducibility hooks.",
    )
    kvp(
        doc,
        "Gap",
        "What the paper does NOT address — open questions, suggested "
        "future experiments.",
    )
    body(
        doc,
        "Treat reviewer feedback as a checklist for revision, not as ground "
        "truth — LLM reviewers are reasonable but not infallible.",
    )

    doc.add_heading("Common issues", level=2)
    kvp(
        doc,
        "Backend health says api_key_configured: no",
        "Edit backend/.env, restart uvicorn.",
    )
    kvp(
        doc,
        "Frontend cannot reach backend",
        "Confirm NEXT_PUBLIC_API_URL in the frontend .env.local matches the "
        "uvicorn URL, and that the backend's CORS_ALLOWED_ORIGINS lists the "
        "frontend URL.",
    )
    kvp(
        doc,
        "Pipeline times out",
        "A full Write run with reviews takes 60–180s. Don't background the "
        "browser tab — Chrome throttles SSE in background tabs.",
    )

    doc.add_page_break()

    # ============================================================
    # PART B — CREATING THE APP
    # ============================================================
    doc.add_heading("Part B — How CorpusAI was built", level=1)

    doc.add_heading("The big picture", level=2)
    body(
        doc,
        "Two halves: a Python backend (FastAPI + Microsoft Agent Framework) "
        "and a Next.js frontend (React + Tailwind). The backend orchestrates "
        "LLM agents into pipelines and streams progress. The frontend is a "
        "thin presentation layer that uploads PDFs, initiates pipelines, and "
        "renders streamed events.",
    )
    body(
        doc,
        "Three pipelines exist, all wired the same way:",
    )
    code_block(
        doc,
        "Explore:  TopicAnalyser -> SemanticScholar -> PaperSummariser (xN, parallel)\n"
        "          -> GapFinder -> IdeaGenerator -> MethodDesigner\n"
        "          -> ExploreDiscussionWriter\n"
        "\n"
        "Write:    PdfReader + Chunker + BM25 index (Phase 3)\n"
        "          -> for each section in {Intro, Methods, Results, Discussion, Abstract}:\n"
        "                retrieve top-k chunks -> SectionWriter\n"
        "          -> ReferenceFormatter -> assemble markdown\n"
        "          -> [BiologyReviewer || StatisticsReviewer || GapReviewer] (parallel)\n"
        "          -> ReviewSynthesiser",
    )

    doc.add_heading("The agent pattern", level=2)
    body(
        doc,
        "Every agent in the project follows the same three-piece recipe:",
    )
    numbered(
        doc,
        "A Pydantic input model and a Pydantic output model — both live "
        "in app/schemas/.",
    )
    numbered(
        doc,
        "A markdown prompt file in app/prompts/ that defines the agent's "
        "personality and the exact JSON shape it must reply with.",
    )
    numbered(
        doc,
        "A small Python class in app/agents/ that subclasses BaseAgent and "
        "declares prompt_name, output_model, and a render_input method.",
    )
    body(
        doc,
        "BaseAgent (in app/agents/base.py) handles the rest: constructing the "
        "Microsoft Agent Framework client, loading the prompt, calling the LLM, "
        "extracting a JSON object from the reply, and validating it against "
        "the output model. A new agent costs about 10 lines of Python plus a "
        "markdown file.",
    )
    code_block(
        doc,
        "# A complete agent — TopicAnalyser\n"
        "from app.agents.base import BaseAgent\n"
        "from app.schemas.research import TopicAnalysis, TopicRequest\n"
        "\n"
        "class TopicAnalyser(BaseAgent[TopicRequest, TopicAnalysis]):\n"
        "    prompt_name = \"topic_analyser\"\n"
        "    output_model = TopicAnalysis\n"
        "\n"
        "    def render_input(self, payload: TopicRequest) -> str:\n"
        "        if payload.sub_field:\n"
        "            return f\"Topic: {payload.topic}\\nSub-field: {payload.sub_field}\"\n"
        "        return f\"Topic: {payload.topic}\"",
    )

    doc.add_heading("Backend layout", level=2)
    code_block(
        doc,
        "backend/app/\n"
        "  core/\n"
        "    config.py            settings (env-driven, pydantic-settings)\n"
        "    prompts.py           load_prompt() helper\n"
        "  agents/                13 .py files — one class per agent\n"
        "  prompts/               18 .md files — one prompt per agent\n"
        "  schemas/\n"
        "    research.py          Explore-mode types\n"
        "    papers.py            ingestion types (PDF, chunks, ref sets)\n"
        "    paper.py             Write-mode types (sections, reviews)\n"
        "  services/\n"
        "    semantic_scholar.py  HTTP client for paper search\n"
        "    pdf_reader.py        PDF parsing (pypdf)\n"
        "    chunker.py           paragraph batching into ~350-word chunks\n"
        "    retrieval.py         BM25 index over chunks\n"
        "    reference_store.py   in-memory store of upload sets\n"
        "    docx_export.py       markdown -> .docx writer\n"
        "  workflows/\n"
        "    explore.py           Explore orchestrator\n"
        "    write.py             Write orchestrator (sections + parallel review)\n"
        "  api/\n"
        "    routes_research.py   POST /api/research/explore (SSE)\n"
        "    routes_papers.py     PDF upload + POST /api/paper/write (SSE)\n"
        "  main.py                FastAPI app + CORS",
    )

    doc.add_heading("Schemas — the contract", level=2)
    body(
        doc,
        "Every model that crosses a boundary (HTTP request, agent input, "
        "agent output, workflow event) is a Pydantic model. Three files split "
        "by domain:",
    )
    kvp(doc, "research.py", "Explore mode — TopicRequest, ResearchGap, ResearchIdea, etc.")
    kvp(doc, "papers.py", "Ingestion — UploadedRef, Chunk, RefSet, SearchRequest.")
    kvp(
        doc,
        "paper.py",
        "Write mode — SectionName enum, WriterInput, BiologyReview, "
        "ReviewSynthesis, PaperDraft, WriteResult.",
    )
    body(
        doc,
        "Pydantic does triple duty: HTTP body validation, automatic OpenAPI "
        "schema generation, and JSON parsing of LLM replies (BaseAgent calls "
        "model_validate_json on every reply).",
    )

    doc.add_heading("Prompts — the personalities", level=2)
    body(
        doc,
        "Every agent has a markdown prompt in app/prompts/. They share a strict "
        "convention: end with a fenced JSON block showing exactly the output "
        "shape. That convention is what makes BaseAgent's parser reliable.",
    )
    code_block(
        doc,
        "Reply with ONE JSON object and nothing else:\n"
        "\n"
        "```json\n"
        "{ \"idea\": \"<one paragraph, 120-200 words>\" }\n"
        "```",
    )
    body(
        doc,
        "Edit prompts, not Python, when you want an agent to behave differently. "
        "Anything you'd put in a long docstring instructing the LLM belongs "
        "in the prompt file.",
    )

    doc.add_heading("Services — non-LLM helpers", level=2)
    kvp(
        doc,
        "semantic_scholar.py",
        "Async HTTPX client hitting Semantic Scholar's free /paper/search "
        "endpoint. Filters out papers without abstracts.",
    )
    kvp(
        doc,
        "pdf_reader.py",
        "Wraps pypdf. Returns a list of PageText with 1-indexed page numbers "
        "so chunk metadata can carry citation hints.",
    )
    kvp(
        doc,
        "chunker.py",
        "Paragraph-batching: split each page on blank lines, accumulate up to "
        "~350 words, emit a chunk with the starting page recorded.",
    )
    kvp(
        doc,
        "retrieval.py",
        "BM25 index using rank-bm25. The public surface is add(chunks) and "
        "search(query, k); swap the internals for embeddings later without "
        "changing callers.",
    )
    kvp(
        doc,
        "reference_store.py",
        "Process-local dict keyed by set_id. Holds metadata plus a ChunkIndex "
        "per uploaded set. Single-process MVP — restarts wipe state.",
    )
    kvp(
        doc,
        "docx_export.py",
        "Line-by-line markdown walker that emits a .docx using python-docx.",
    )

    doc.add_heading("Workflows — the orchestrators", level=2)
    body(
        doc,
        "Per the project's rules, agents do not call each other directly. "
        "Both Explore and Write are async generators that yield typed events. "
        "Each yield becomes one Server-Sent Event the browser consumes.",
    )
    code_block(
        doc,
        "@dataclass\n"
        "class ExploreEvent:\n"
        "    kind: EventKind         # 'started', 'idea_generated', ...\n"
        "    payload: dict[str, Any]\n"
        "\n"
        "async def run_explore(request) -> AsyncIterator[ExploreEvent]:\n"
        "    yield ExploreEvent('started', {...})\n"
        "    analysis = await TopicAnalyser().run(request)\n"
        "    yield ExploreEvent('topic_analysed', {...})\n"
        "    # ...etc",
    )
    body(
        doc,
        "Two patterns worth noting:",
    )
    bullet(
        doc,
        "Concurrency where it pays. PaperSummariser runs over N papers via "
        "asyncio.gather — total wait is the slowest paper, not the sum. "
        "Same trick for the three reviewers in Write mode.",
    )
    bullet(
        doc,
        "Dependency injection for tests. run_explore takes an optional "
        "search_fn parameter; run_write takes an optional store. Tests pass "
        "fakes; production uses the real services.",
    )

    doc.add_heading("API — thin pass-throughs", level=2)
    body(
        doc,
        "Routes do almost no logic — they validate the request, call the "
        "workflow, and stream events. Errors during the stream are emitted "
        "as 'error' events rather than HTTP 500s, so the client always sees "
        "a clean stream end.",
    )
    code_block(
        doc,
        "@router.post(\"/explore\")\n"
        "async def explore(request: TopicRequest) -> EventSourceResponse:\n"
        "    return EventSourceResponse(_event_stream(request))\n"
        "\n"
        "async def _event_stream(req):\n"
        "    try:\n"
        "        async for event in run_explore(req):\n"
        "            yield {\"event\": event.kind, \"data\": json.dumps(event.payload)}\n"
        "    except Exception as exc:\n"
        "        yield {\"event\": \"error\", \"data\": json.dumps({\"message\": str(exc)})}",
    )

    doc.add_heading("Frontend layout", level=2)
    code_block(
        doc,
        "frontend/\n"
        "  app/\n"
        "    layout.tsx           html shell + fonts\n"
        "    page.tsx             header/footer + ModeTabs\n"
        "  components/\n"
        "    ModeTabs.tsx         Explore | Write tab switcher\n"
        "    explore/\n"
        "      ExploreView.tsx    sidebar + streaming result panel\n"
        "    write/\n"
        "      WriteView.tsx      whole Write workflow UI\n"
        "      PdfUploader.tsx    multi-file <input type=file>\n"
        "      SectionSelector.tsx  five checkboxes\n"
        "      ReviewPanel.tsx    four-tab review viewer\n"
        "      ExportButtons.tsx  .docx (backend) + .md (Blob)\n"
        "    ui/\n"
        "      StreamLog.tsx      generic step list\n"
        "      Markdown.tsx       react-markdown wrapper\n"
        "  lib/\n"
        "    types.ts             TS mirrors of Pydantic schemas\n"
        "    api.ts               fetch wrappers (upload, exportDocx)\n"
        "    sse.ts               manual SSE parser over fetch streams",
    )

    doc.add_heading("The SSE pattern (most important frontend file)", level=2)
    body(
        doc,
        "Browser EventSource only supports GET, but our streams are POST + "
        "JSON body. Hand-rolled async generator over fetch's ReadableStream:",
    )
    code_block(
        doc,
        "// lib/sse.ts (excerpt)\n"
        "export async function* streamSse(url, body) {\n"
        "  const response = await fetch(url, {\n"
        "    method: 'POST',\n"
        "    headers: {'Content-Type': 'application/json'},\n"
        "    body: JSON.stringify(body),\n"
        "  });\n"
        "  const reader = response.body.getReader();\n"
        "  const decoder = new TextDecoder();\n"
        "  let buffer = '';\n"
        "  while (true) {\n"
        "    const {value, done} = await reader.read();\n"
        "    if (done) break;\n"
        "    buffer += decoder.decode(value, {stream: true});\n"
        "    // split on \\n\\n, parse 'event:' and 'data:' lines, yield {kind,data}\n"
        "  }\n"
        "}",
    )
    body(
        doc,
        "Both ExploreView and WriteView consume this generator with for-await:",
    )
    code_block(
        doc,
        "for await (const ev of streamSse(url, payload)) {\n"
        "  if (ev.kind === 'section_completed') updateProgress(ev.data.section);\n"
        "  if (ev.kind === 'completed') setResult(ev.data);\n"
        "}",
    )

    doc.add_heading("Recipe — adding a new agent", level=2)
    body(
        doc,
        "Say you want to add a CitationChecker that verifies references "
        "against the literature.",
    )
    numbered(
        doc,
        "Add the input/output models in app/schemas/paper.py "
        "(CitationCheckerInput, CitationCheckResult).",
    )
    numbered(
        doc,
        "Write app/prompts/citation_checker.md with the agent's instructions "
        "and the JSON output spec.",
    )
    numbered(
        doc,
        "Write app/agents/citation_checker.py:",
    )
    code_block(
        doc,
        "from app.agents.base import BaseAgent\n"
        "from app.schemas.paper import CitationCheckerInput, CitationCheckResult\n"
        "\n"
        "class CitationChecker(BaseAgent[CitationCheckerInput, CitationCheckResult]):\n"
        "    prompt_name = \"citation_checker\"\n"
        "    output_model = CitationCheckResult\n"
        "\n"
        "    def render_input(self, payload):\n"
        "        return f\"Paper:\\n{payload.markdown}\"",
    )
    numbered(doc, "Export the class from app/agents/__init__.py.")
    numbered(
        doc,
        "Call it from the relevant workflow (e.g. between paper_assembled "
        "and review_started). Yield a new event kind for the UI.",
    )
    numbered(
        doc,
        "Write a test in tests/test_agents/ that monkey-patches the agent's "
        "_agent.run and verifies the parser populates the model.",
    )

    doc.add_heading("Tests", level=2)
    body(
        doc,
        "Every test runs without an Anthropic API key. The trick is to "
        "monkey-patch BaseAgent's _agent.run with a fake that returns canned "
        "JSON. PDFs in tests are generated in-process by reportlab so no "
        "fixture files are committed.",
    )
    code_block(
        doc,
        "uv run pytest -v\n"
        "# 26 passed in ~3s",
    )

    doc.add_heading("Deployment summary", level=2)
    body(
        doc,
        "Backend on Fly.io (single always-on machine — required because the "
        "ReferenceStore is in-process). Frontend on Vercel (free tier, "
        "auto-detected Next.js). See DEPLOY.md at the repo root for the "
        "full step-by-step.",
    )

    doc.add_heading("Where to read next", level=2)
    bullet(doc, "CLAUDE.md — project conventions and rules.")
    bullet(doc, ".claude/rules/agent-conventions.md — what every agent must do.")
    bullet(doc, "DEPLOY.md — Fly.io + Vercel walkthrough.")

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    main()
